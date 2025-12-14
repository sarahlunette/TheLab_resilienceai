"""
Qdrant Vectorstore Builder (v6) - Fixed for large documents & batching
- pdfplumber for PDF extraction
- DOCX, TXT, MD, CSV support
- NEW: JSON ingestion (arrays, dicts, nested)
- LOCAL embedding model: ../models/all-MiniLM-L6-v2
- Automatic sync with Qdrant:
      ✔ add new files
      ✔ remove deleted files
"""

from pathlib import Path
import pandas as pd
import pdfplumber
import json
import os
from dotenv import load_dotenv

from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance, PointStruct

from llama_index.core import Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

load_dotenv()

# ============================================================
# CONFIG
# ============================================================
DOCS_DIR = Path("./docs")
QDRANT_URL = os.environ["QDRANT_URL"]
QDRANT_API_KEY = os.environ["QDRANT_API_KEY"]
COLLECTION_NAME = "island_docs"
EMBEDDING_MODEL = "./models/all-MiniLM-L6-v2"

CHUNK_SIZE = 500  # words per chunk
BATCH_SIZE = 200  # points per upsert


# ============================================================
# HELPERS — Text Chunking
# ============================================================
def chunk_text(text, chunk_size=CHUNK_SIZE):
    words = text.split()
    for i in range(0, len(words), chunk_size):
        yield " ".join(words[i : i + chunk_size])


# ============================================================
# HELPERS — Loaders (unchanged)
# ============================================================
def load_pdf(path: Path) -> list[Document]:
    docs = []
    try:
        with pdfplumber.open(path) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
        if text.strip():
            docs.append(Document(text=text, metadata={"source": path.name}))
        else:
            print(f"⚠ PDF EMPTY: {path.name}")
    except Exception as e:
        print(f"❌ Error reading PDF {path.name}: {e}")
    return docs


def load_csv_rows(path: Path) -> list[Document]:
    df = pd.read_csv(path)
    docs = []
    for i, row in df.iterrows():
        text = "\n".join(f"{col}: {row[col]}" for col in df.columns)
        if text.strip():
            docs.append(
                Document(text=text, metadata={"source": path.name, "row_index": int(i)})
            )
    return docs


def load_text_file(path: Path) -> list[Document]:
    try:
        text = path.read_text(errors="ignore")
        if text.strip():
            return [Document(text=text, metadata={"source": path.name})]
    except:
        pass
    return []


def load_docx(path: Path) -> list[Document]:
    try:
        import docx

        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        return [Document(text=text, metadata={"source": path.name})]
    except:
        print(f"❌ Cannot read DOCX: {path.name}")
        return []


def load_json(path: Path) -> list[Document]:
    try:
        data = json.loads(path.read_text(errors="ignore"))
    except Exception as e:
        print(f"❌ Error reading JSON {path.name}: {e}")
        return []

    docs = []

    def flatten(obj, prefix=""):
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                key = f"{prefix}{k}"
                if isinstance(v, (dict, list)):
                    lines.append(f"{key}:")
                    lines.append(flatten(v, prefix=prefix + "  "))
                else:
                    lines.append(f"{key}: {v}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                lines.append(f"{prefix}- item {i}:")
                lines.append(flatten(item, prefix=prefix + "  "))
        else:
            lines.append(f"{prefix}{obj}")
        return "\n".join(lines)

    if isinstance(data, list):
        for idx, item in enumerate(data):
            text = flatten(item)
            docs.append(
                Document(text=text, metadata={"source": path.name, "json_index": idx})
            )
    elif isinstance(data, dict):
        text = flatten(data)
        docs.append(Document(text=text, metadata={"source": path.name}))
    else:
        print(f"⚠ Unsupported JSON structure in {path.name}")

    return docs


# ============================================================
# KEY GENERATOR
# ============================================================
def doc_key(doc: Document) -> str:
    key = doc.metadata.get("source")
    for suffix in ("row_index", "json_index"):
        if suffix in doc.metadata:
            key += f"#{suffix}:{doc.metadata[suffix]}"
    return key


# ============================================================
# UPLOAD POINTS IN BATCH
# ============================================================
def upsert_points_in_batches(client, collection_name, points):
    for i in range(0, len(points), BATCH_SIZE):
        batch = points[i : i + BATCH_SIZE]
        client.upsert(collection_name=collection_name, points=batch)


# ============================================================
# SYNC FUNCTION — Add new docs / remove deleted ones
# ============================================================
def sync_vectorstore(client, collection_name, embed_model, docs):
    print("\n🔄 Syncing Qdrant collection…")

    try:
        existing, _ = client.scroll(collection_name, limit=50_000)
    except Exception as e:
        print(f"❌ Could not scroll existing vectors: {e}")
        return

    existing_index = {}
    for point in existing:
        key = point.payload.get("source", "")
        for suffix in ("row_index", "json_index"):
            if suffix in point.payload:
                key += f"#{suffix}:{point.payload[suffix]}"
        existing_index[key] = point.id

    new_index = {doc_key(d): d for d in docs}

    # Deletions
    to_delete = [existing_index[k] for k in existing_index if k not in new_index]
    if to_delete:
        print(f"🗑 Removing {len(to_delete)} removed docs…")
        client.delete(collection_name, points_selector={"points": to_delete})
    else:
        print("🟢 No deletions needed.")

    # Additions
    to_add = [new_index[k] for k in new_index if k not in existing_index]
    if to_add:
        print(f"➕ Adding {len(to_add)} new docs…")

        next_id = max(existing_index.values(), default=0) + 1
        new_points = []

        for doc in to_add:
            for chunk in chunk_text(doc.text, chunk_size=CHUNK_SIZE):
                vector = embed_model.get_text_embedding(chunk)
                new_points.append(
                    PointStruct(
                        id=next_id,
                        vector=vector,
                        payload={"text": chunk, **doc.metadata},
                    )
                )
                next_id += 1

        upsert_points_in_batches(client, collection_name, new_points)
    else:
        print("🟢 No new additions.")

    print("✔ Sync complete.")


# ============================================================
# MAIN
# ============================================================
def main():
    print("\n=== 🚀 QDRANT VECTORSTORE BUILDER v6 (Fixed) ===")

    if not DOCS_DIR.exists():
        raise FileNotFoundError("docs/ folder missing")

    # Load documents
    all_docs = []
    print("📚 Loading documents…")
    for file in DOCS_DIR.iterdir():
        suffix = file.suffix.lower()
        if suffix == ".pdf":
            all_docs.extend(load_pdf(file))
        elif suffix in (".txt", ".md"):
            all_docs.extend(load_text_file(file))
        elif suffix == ".docx":
            all_docs.extend(load_docx(file))
        elif suffix == ".csv":
            all_docs.extend(load_csv_rows(file))
        elif suffix == ".json":
            all_docs.extend(load_json(file))
    print(f"📦 Loaded documents: {len(all_docs)}")

    if not all_docs:
        print("⚠ No documents — stopping.")
        return

    # Embedding model
    print(f"🧠 Loading embedding model: {EMBEDDING_MODEL}")
    embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
    embedding_dim = len(embed_model.get_text_embedding("test"))
    print(f"🔢 Embedding size = {embedding_dim}")

    # Qdrant connection
    print("🗄 Connecting to Qdrant…")
    client = QdrantClient(
        url=QDRANT_URL, api_key=QDRANT_API_KEY, prefer_grpc=False, timeout=60
    )

    collections = [c.name for c in client.get_collections().collections]

    # Create or Sync
    if COLLECTION_NAME not in collections:
        print(f"📌 Creating new collection `{COLLECTION_NAME}`")
        client.recreate_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=embedding_dim, distance=Distance.COSINE),
        )

        points = []
        pid = 1
        for doc in all_docs:
            for chunk in chunk_text(doc.text, chunk_size=CHUNK_SIZE):
                vector = embed_model.get_text_embedding(chunk)
                points.append(
                    PointStruct(
                        id=pid, vector=vector, payload={"text": chunk, **doc.metadata}
                    )
                )
                pid += 1

        upsert_points_in_batches(client, COLLECTION_NAME, points)
        print("✔ Full ingestion complete.")
    else:
        print(f"🔄 Collection exists — syncing…")
        sync_vectorstore(client, COLLECTION_NAME, embed_model, all_docs)

    # Stats
    count = client.count(COLLECTION_NAME).count
    print("\n🎉 DONE — Qdrant Vectorstore Ready!")
    print(f"📌 Collection: {COLLECTION_NAME}")
    print(f"📊 Total vectors stored: {count}")
    print("🚀 Ready for RAG pipelines.\n")


if __name__ == "__main__":
    main()
